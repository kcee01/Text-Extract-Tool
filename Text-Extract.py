import os
from PyPDF2 import PdfReader
from docx import Document
from tkinter import Tk, filedialog

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file page by page, preserving structure."""
    reader = PdfReader(pdf_path)
    all_pages = []

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        # Add a clear page separator to preserve structure
        page_marker = f"\n\n--- Page {i} of {len(reader.pages)} ---\n\n"
        all_pages.append(page_marker + text)

    return "\n".join(all_pages).strip()

def save_to_txt(text, output_path):
    """Save extracted text to a .txt file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

def save_to_docx(text, output_path):
    """Save extracted text to a Word .docx file, preserving page structure."""
    doc = Document()
    sections = text.split("\n\n--- Page")
    
    for idx, section in enumerate(sections):
        if not section.strip():
            continue
        if idx == 0:
            # First section (before first page marker)
            doc.add_paragraph(section.strip())
        else:
            # Re-add the "Page" label and insert a page break before it
            doc.add_page_break()
            doc.add_paragraph(f"--- Page{section}")

    doc.save(output_path)

def pdf_to_text_or_word(pdf_path, output_format="txt"):
    """Extract text from a PDF and save as either .txt or .docx."""
    if not os.path.exists(pdf_path):
        raise FileNotFoundError("PDF file not found.")
    
    text = extract_text_from_pdf(pdf_path)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    if output_format == "txt":
        output_path = os.path.join(script_dir, base_name + ".txt")
        save_to_txt(text, output_path)
    elif output_format == "docx":
        output_path = os.path.join(script_dir, base_name + ".docx")
        save_to_docx(text, output_path)
    else:
        raise ValueError("Output format must be 'txt' or 'docx'.")

    print(f"\n✅ File successfully saved at: {output_path}")

# -------------------------------
# Interactive Section (with File Picker)
# -------------------------------
if __name__ == "__main__":
    print("📄 PDF to Text/Word Converter (Preserves Page Order)")
    print("----------------------------------------------------")

    # Open file picker dialog
    root = Tk()
    root.withdraw()  # Hide the main Tk window
    pdf_path = filedialog.askopenfilename(
        title="Select PDF file to convert",
        filetypes=[("PDF files", "*.pdf")],
    )

    if not pdf_path:
        print("❌ No file selected. Exiting...")
        exit()

    print(f"📂 Selected file: {pdf_path}")

    # Ask for output format
    output_format = input("Choose output format ('txt' or 'docx') [default=txt]: ").strip().lower()
    if output_format not in ["txt", "docx", ""]:
        print("⚠️ Invalid choice. Defaulting to 'txt'.")
        output_format = "txt"
    elif output_format == "":
        output_format = "txt"

    try:
        pdf_to_text_or_word(pdf_path, output_format)
    except Exception as e:
        print(f"❌ Error: {e}")
